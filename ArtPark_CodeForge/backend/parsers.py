"""
parsers.py — Document text extraction utilities.
Supports PDF (pdfplumber), DOCX (python-docx), and plain text.
"""

import io
import re
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                # Also extract text from tables
                for table in page.extract_tables():
                    for row in table:
                        row_text = " | ".join(cell or "" for cell in row)
                        if row_text.strip():
                            text_parts.append(row_text)
        raw = "\n".join(text_parts)
        return _clean_text(raw)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        # Fallback to PyPDF2
        return _extract_pdf_fallback(file_bytes)


def _extract_pdf_fallback(file_bytes: bytes) -> str:
    """Fallback PDF extraction using PyPDF2."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return _clean_text("\n".join(pages))
    except Exception as e:
        logger.error(f"PyPDF2 fallback failed: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return _clean_text("\n".join(parts))
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode plain text with encoding detection."""
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            return _clean_text(file_bytes.decode(encoding))
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct extractor based on file extension."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        # Best-effort: try docx parser (may partially work)
        logger.warning("Legacy .doc format — partial extraction only")
        return extract_text_from_docx(file_bytes)
    else:
        # txt, md, etc.
        return extract_text_from_txt(file_bytes)


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove junk characters."""
    # Remove null bytes and non-printable chars (keep newlines/tabs)
    text = re.sub(r"[^\x09\x0a\x0d\x20-\x7e\u00a0-\uffff]", " ", text)
    # Collapse multiple spaces / tabs
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Collapse 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
